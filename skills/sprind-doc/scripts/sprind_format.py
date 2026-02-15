#!/usr/bin/env python3
"""
SPRIND Document Formatter

Creates SPRIND-branded .docx documents using the official SPRIND template as base.
The template provides:
- Full-page background images (logo/letterhead) in headers
- GT America fonts with proper SPRIND paragraph styles
- Correct page layout (A4, margins, header/footer distances)

This script adds:
- Parsed content mapped to SPRIND styles
- Version/date footer with page numbers
- Auto-detected language (EN/DE) for footer format

Usage:
    python3 sprind_format.py --input content.md --output document.docx
    python3 sprind_format.py --input content.md --output document.docx --version 1.0 --date "15.02.2025"
    python3 sprind_format.py --input existing.docx --output reformatted.docx
    echo "Some text" | python3 sprind_format.py --output document.docx
"""

import argparse
import os
import re
import shutil
import sys
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips, RGBColor

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

ASSETS_DIR = Path(__file__).resolve().parent.parent / "assets"
TEMPLATE_DOTX = ASSETS_DIR / "SPRIND_Vorlage_GTAmerica.dotx"

# SPRIND style names (as defined in the template, with proper Umlauts)
STYLE_HEADING = "SPRIND - Überschrift"
STYLE_SUBHEADING = "SPRIND - Unterüberschrift"
STYLE_PARAGRAPH = "SPRIND - Paragraph"
STYLE_BULLET = "SPRIND - Auflistung"
STYLE_NUMBERED = "SPRIND - Aufzählung"
STYLE_EMPHASIS = "SPRIND - Hervorhebung"
STYLE_FOOTNOTE = "SPRIND - Fußnote"
STYLE_NUMBERED_SUBHEADING = "SPRIND - Nummerierte Unterüberschrift"

FONT_BODY = "GT America Light"

# German detection words
GERMAN_MARKERS = {
    "der", "die", "das", "und", "oder", "für", "über", "aber", "nach",
    "mit", "von", "bei", "seit", "wird", "werden", "haben", "sein",
    "einen", "einer", "eines", "einem", "nicht", "auch", "sich",
    "dass", "diese", "dieser", "dieses", "diesem", "können", "müssen",
    "sollen", "zwischen", "durch", "bereits", "sowie", "jedoch",
    "gegenüber", "während", "innerhalb", "außerdem", "grundsätzlich",
}


# ---------------------------------------------------------------------------
# Template loading
# ---------------------------------------------------------------------------

def load_template() -> Document:
    """Load the SPRIND .dotx template as a .docx Document.

    python-docx refuses to open .dotx files directly, so we copy to a
    temporary .docx and patch the content type inside the ZIP.
    """
    if not TEMPLATE_DOTX.exists():
        print(f"Error: Template not found: {TEMPLATE_DOTX}", file=sys.stderr)
        sys.exit(1)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()

    # Read the .dotx, swap content type to .docx, write out
    with zipfile.ZipFile(str(TEMPLATE_DOTX), "r") as zin:
        with zipfile.ZipFile(tmp.name, "w", zipfile.ZIP_DEFLATED) as zout:
            for name in zin.namelist():
                data = zin.read(name)
                if name == "[Content_Types].xml":
                    data = data.replace(
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    )
                zout.writestr(name, data)

    doc = Document(tmp.name)
    os.unlink(tmp.name)
    return doc


def clear_body(doc: Document):
    """Remove all body paragraphs from the document, preserving section properties."""
    body = doc.element.body
    for p in list(body.findall(qn("w:p"))):
        # Keep the paragraph that contains sectPr (section properties)
        if p.find(f"{qn('w:pPr')}/{qn('w:sectPr')}") is not None:
            continue
        body.remove(p)
    # Also remove any tables from template
    for tbl in list(body.findall(qn("w:tbl"))):
        body.remove(tbl)


# ---------------------------------------------------------------------------
# Language detection
# ---------------------------------------------------------------------------

def detect_language(text: str) -> str:
    """Detect whether text is primarily German or English."""
    words = re.findall(r"\b\w+\b", text.lower())
    if not words:
        return "en"
    german_count = sum(1 for w in words if w in GERMAN_MARKERS)
    ratio = german_count / len(words)
    return "de" if ratio > 0.03 else "en"


# ---------------------------------------------------------------------------
# Date formatting
# ---------------------------------------------------------------------------

def format_date_english(date_str: str | None) -> str:
    """Format date for English footer: '19th March 2024'."""
    dt = _parse_date(date_str)
    day = dt.day
    suffix = {1: "st", 2: "nd", 3: "rd", 21: "st", 22: "nd", 23: "rd", 31: "st"}.get(day, "th")
    return f"{day}{suffix} {dt.strftime('%B %Y')}"


def format_date_german(date_str: str | None) -> str:
    """Format date for German footer: '15.02.2025'."""
    dt = _parse_date(date_str)
    return dt.strftime("%d.%m.%Y")


def _parse_date(date_str: str | None) -> datetime:
    """Parse a date string in various formats, default to today."""
    if not date_str:
        return datetime.now()
    for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y", "%B %d, %Y", "%d %B %Y"):
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    return datetime.now()


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

class ContentBlock:
    """Represents a parsed content block."""
    def __init__(self, kind: str, text: str, level: int = 0, table_data: list = None):
        self.kind = kind      # "heading", "paragraph", "bullet", "numbered", "table"
        self.text = text
        self.level = level    # heading level (1, 2, 3...)
        self.table_data = table_data  # list of rows, each row is list of cell texts


def parse_markdown(text: str) -> list[ContentBlock]:
    """Parse markdown text into ContentBlock list."""
    blocks = []
    lines = text.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            blocks.append(ContentBlock("heading", m.group(2).strip(), level=level))
            i += 1
            continue

        # Bullet list
        m = re.match(r"^[-*+]\s+(.+)$", stripped)
        if m:
            blocks.append(ContentBlock("bullet", m.group(1).strip()))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if m:
            blocks.append(ContentBlock("numbered", m.group(1).strip()))
            i += 1
            continue

        # Regular paragraph (collect consecutive non-empty lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6}\s|[-*+]\s|\d+[.)]\s)", lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append(ContentBlock("paragraph", " ".join(para_lines)))

    return blocks


def parse_inline(text: str) -> list[tuple]:
    """Parse inline markdown (bold, italic, underline) into runs.
    Returns list of (text, bold, italic, underline) tuples."""
    runs = []
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|__(.+?)__|([^*_]+))"
    for m in re.finditer(pattern, text):
        if m.group(2):      # **bold**
            runs.append((m.group(2), True, False, False))
        elif m.group(3):    # *italic*
            runs.append((m.group(3), False, True, False))
        elif m.group(4):    # __underline__
            runs.append((m.group(4), False, False, True))
        elif m.group(5):    # plain text
            runs.append((m.group(5), False, False, False))
    if not runs:
        runs.append((text, False, False, False))
    return runs


def _is_heading_by_format(para) -> int | None:
    """Detect if a Normal-styled paragraph is actually a heading based on font size.
    Returns heading level (1 or 2) or None."""
    if not para.runs:
        return None
    # Check the largest font size across runs
    max_size = 0
    for run in para.runs:
        if run.font.size:
            max_size = max(max_size, run.font.size)
    # 16pt+ = H1 (Überschrift), 13pt+ = H2 (Unterüberschrift)
    if max_size >= 200000:   # ~16pt in EMU
        return 1
    elif max_size >= 165000:  # ~13pt in EMU
        return 2
    return None


def parse_docx_input(filepath: str) -> list[ContentBlock]:
    """Extract content blocks from an existing .docx file, preserving tables."""
    doc = Document(filepath)
    blocks = []

    # Build an ordered list of body elements (paragraphs and tables interleaved)
    body = doc.element.body
    para_index = 0
    table_index = 0
    paragraphs = doc.paragraphs
    tables = doc.tables

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            if para_index < len(paragraphs):
                para = paragraphs[para_index]
                para_index += 1
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name.lower() if para.style else ""

                # Check explicit heading styles
                if "heading" in style_name or "überschrift" in style_name:
                    level = 1
                    if "unter" in style_name or "zwischen" in style_name or "2" in style_name:
                        level = 2
                    elif "3" in style_name:
                        level = 3
                    blocks.append(ContentBlock("heading", text, level=level))
                # Check if Normal paragraph looks like a heading (large font)
                elif style_name in ("normal", "body text", "default paragraph font", ""):
                    heading_level = _is_heading_by_format(para)
                    if heading_level:
                        blocks.append(ContentBlock("heading", text, level=heading_level))
                    else:
                        blocks.append(ContentBlock("paragraph", text))
                elif "list" in style_name or "bullet" in style_name or "auflistung" in style_name:
                    blocks.append(ContentBlock("bullet", text))
                elif "number" in style_name or "aufzählung" in style_name:
                    blocks.append(ContentBlock("numbered", text))
                else:
                    blocks.append(ContentBlock("paragraph", text))

        elif tag == "tbl":
            if table_index < len(tables):
                table = tables[table_index]
                table_index += 1
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                if rows:
                    blocks.append(ContentBlock("table", "", table_data=rows))

    return blocks


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def create_sprind_document(
    blocks: list[ContentBlock],
    output_path: str,
    version: str | None = None,
    date_str: str | None = None,
    language: str | None = None,
    title: str | None = None,
):
    """Create a SPRIND-formatted .docx from content blocks."""

    # Auto-detect language from content if not specified
    all_text = " ".join(b.text for b in blocks)
    if not language:
        language = detect_language(all_text)

    # Load template and clear placeholder content
    doc = load_template()
    clear_body(doc)

    # Update footers
    section = doc.sections[0]
    _update_footer(section, version=version, date_str=date_str, language=language, is_first_page=True)
    _update_footer(section, version=version, date_str=date_str, language=language, is_first_page=False)

    # Add content
    for block in blocks:
        if block.kind == "heading":
            if block.level <= 1:
                p = doc.add_paragraph(style=STYLE_HEADING)
            else:
                p = doc.add_paragraph(style=STYLE_SUBHEADING)
            _add_runs(p, block.text)

        elif block.kind == "bullet":
            p = doc.add_paragraph(style=STYLE_BULLET)
            _add_runs(p, block.text)

        elif block.kind == "numbered":
            p = doc.add_paragraph(style=STYLE_NUMBERED)
            _add_runs(p, block.text)

        elif block.kind == "table":
            _add_table(doc, block.table_data)

        elif block.kind == "paragraph":
            p = doc.add_paragraph(style=STYLE_PARAGRAPH)
            _add_runs(p, block.text)

    doc.save(output_path)
    print(f"Created: {output_path}")


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

def _update_footer(section, version: str | None, date_str: str | None, language: str, is_first_page: bool):
    """Replace footer content with version/date and page numbers."""
    if is_first_page:
        footer = section.first_page_footer
        font_size = Pt(10) if language == "en" else Pt(8)
    else:
        footer = section.footer
        font_size = Pt(10) if language == "en" else Pt(8)

    footer.is_linked_to_previous = False

    # Clear existing footer content
    for p in list(footer.paragraphs):
        for child in list(p._element):
            p._element.remove(child)

    # Use first existing paragraph or add one
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_before = Twips(360)

    if language == "en":
        # "Version X.Y from DDth Month YYYY, Page N (T)"
        prefix = ""
        if version:
            date_formatted = format_date_english(date_str)
            prefix = f"Version {version} from {date_formatted}, Page "
        else:
            prefix = "Page "
        run = para.add_run(prefix)
        _set_run_font(run, FONT_BODY, font_size)
        _add_field(para, "PAGE", FONT_BODY, font_size)
        run2 = para.add_run(" (")
        _set_run_font(run2, FONT_BODY, font_size)
        _add_field(para, "NUMPAGES", FONT_BODY, font_size)
        run3 = para.add_run(")")
        _set_run_font(run3, FONT_BODY, font_size)
    else:  # German
        # "Version X.Y vom DD.MM.YYYY   Seite N/T"
        date_formatted = format_date_german(date_str)
        if version:
            prefix = f"Version {version} vom {date_formatted}   Seite "
        else:
            prefix = f"{date_formatted}   Seite "
        run = para.add_run(prefix)
        _set_run_font(run, FONT_BODY, font_size)
        _add_field(para, "PAGE", FONT_BODY, font_size)
        run2 = para.add_run("/")
        _set_run_font(run2, FONT_BODY, font_size)
        _add_field(para, "SECTIONPAGES", FONT_BODY, font_size)


def _add_field(paragraph, field_name: str, font_name: str, font_size):
    """Add a Word field code (PAGE, NUMPAGES, SECTIONPAGES) to a paragraph."""
    run = paragraph.add_run()
    _set_run_font(run, font_name, font_size)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    _set_run_font(run2, font_name, font_size)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    _set_run_font(run3, font_name, font_size)
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_char_separate)

    run4 = paragraph.add_run("1")
    _set_run_font(run4, font_name, font_size)

    run5 = paragraph.add_run()
    _set_run_font(run5, font_name, font_size)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run5._element.append(fld_char_end)


def _set_run_font(run, font_name: str, font_size):
    """Set font name and size on a run."""
    run.font.name = font_name
    run.font.size = font_size
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------

def _add_table(doc, table_data: list):
    """Add a table to the document with SPRIND body font styling."""
    if not table_data:
        return
    rows = len(table_data)
    cols = len(table_data[0]) if table_data else 0
    if cols == 0:
        return

    table = doc.add_table(rows=rows, cols=cols)

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                # Clear default paragraph and add styled one
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.name = FONT_BODY
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
                # Set hAnsi font
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rfonts)
                rfonts.set(qn("w:ascii"), FONT_BODY)
                rfonts.set(qn("w:hAnsi"), FONT_BODY)
                # Bold first row (header) or first column (key-value tables)
                if i == 0 or (cols == 2 and j == 0):
                    run.bold = True


def _add_runs(paragraph, text: str):
    """Add inline-formatted runs to a paragraph."""
    runs = parse_inline(text)
    for text_chunk, bold, italic, underline in runs:
        run = paragraph.add_run(text_chunk)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if underline:
            run.underline = True


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Create SPRIND-formatted .docx documents")
    parser.add_argument("--input", "-i", help="Input file path (.md, .txt, or .docx)")
    parser.add_argument("--output", "-o", required=True, help="Output .docx file path")
    parser.add_argument("--title", "-t", help="Document title")
    parser.add_argument("--version", "-v", help="Version string (e.g., '1.0')")
    parser.add_argument("--date", "-d", help="Date string (default: today)")
    parser.add_argument("--language", "-l", choices=["en", "de"], help="Language (auto-detected if omitted)")
    args = parser.parse_args()

    # Read input
    if args.input:
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"Error: Input file not found: {args.input}", file=sys.stderr)
            sys.exit(1)

        if input_path.suffix.lower() == ".docx":
            blocks = parse_docx_input(str(input_path))
        else:
            with open(input_path, "r", encoding="utf-8") as f:
                content = f.read()
            blocks = parse_markdown(content)
    elif not sys.stdin.isatty():
        content = sys.stdin.read()
        blocks = parse_markdown(content)
    else:
        print("Error: No input provided. Use --input or pipe content via stdin.", file=sys.stderr)
        sys.exit(1)

    # Prepend title if specified and not already in content
    if args.title:
        has_title = any(b.kind == "heading" and b.level <= 1 for b in blocks)
        if not has_title:
            blocks.insert(0, ContentBlock("heading", args.title, level=1))

    create_sprind_document(
        blocks=blocks,
        output_path=args.output,
        version=args.version,
        date_str=args.date,
        language=args.language,
        title=args.title,
    )


if __name__ == "__main__":
    main()
